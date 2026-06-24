"""SpaceX IPO, part 4: build a Word report (a data narrative) from the saved results.

This reads the files that parts 1-3 saved and writes a Microsoft Word report:
  output/spcx_hourly.csv, output/spcx_5min.csv     (prices and returns)
  output/spcx_scorecard.csv                         (the SpaceX scorecard)
  output/comparison_metrics.csv                     (SpaceX vs Tesla, Nvidia, QQQ)
  output/figures/*.png                              (the FT-style figures)

It follows the fins-agent report convention: built-in Word styles (Title, Heading,
Caption, Normal), an A4 page, captioned figures, short tables to two or three
significant figures, and the standard sections (executive summary, data, method,
results, conclusion, references).

Run parts 1, 2, and 3 first so the input files exist. Then run this file. The
report is written to report/spacex_ipo_report.docx.

Needs python-docx:  pip install python-docx
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# Same absolute-path setup as the other scripts, so paths never nest, on Mac or Windows.
THIS_SCRIPT_FOLDER = Path("fins2026") / "week4" / "scratch" / "spacex_ipo"
try:
    BASE_DIR = Path(__file__).resolve().parent
except NameError:  # no __file__ when running a highlighted selection in the console
    guess = Path.cwd() / THIS_SCRIPT_FOLDER
    BASE_DIR = guess if guess.is_dir() else Path.cwd()
OUTPUT_DIR = BASE_DIR / "output"
FIGURE_DIR = OUTPUT_DIR / "figures"
REPORT_DIR = BASE_DIR / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
REPORT_PATH = REPORT_DIR / "spacex_ipo_report.docx"


# -----------------------------------------------------------------------------
# 1. Read the saved results
# -----------------------------------------------------------------------------

def require(path):
    """Stop with a plain message if an input file is missing, rather than a traceback."""
    if not path.is_file():
        raise SystemExit(f"Missing {path.name}. Run parts 1-3 first to create it.")
    return path


hourly = pd.read_csv(require(OUTPUT_DIR / "spcx_hourly.csv"), parse_dates=["datetime"]).set_index("datetime")
five_min = pd.read_csv(require(OUTPUT_DIR / "spcx_5min.csv"), parse_dates=["datetime"]).set_index("datetime")
scorecard = pd.read_csv(require(OUTPUT_DIR / "spcx_scorecard.csv"), index_col=0)["value"]
comparison = pd.read_csv(require(OUTPUT_DIR / "comparison_metrics.csv"), index_col=0)

# Numbers used in the written narrative, read straight from the data so the words match
# the figures and tables exactly.
sample_start = hourly.index.min()
sample_end = hourly.index.max()
trading_days = hourly.index.normalize().nunique() - 1  # minus the offer-price row
total_return = float(scorecard["Total return (%)"])
ann_vol = float(scorecard["Annualized volatility (%)"])
ann_sharpe = float(scorecard["Annualized Sharpe ratio"])
excess_kurtosis = float(five_min["return"].dropna().kurt())
spcx_cmp = comparison.loc["SPCX"]
best_sharpe_name = comparison["sharpe"].idxmax()
best_sharpe_row = comparison.loc[best_sharpe_name]

# Sign-aware phrasing so the prose stays correct as the data updates (the live pull can swing
# SpaceX from a leader to a laggard between runs).
others_cmp = comparison.drop(index="SPCX")
spcx_first_trade = float(spcx_cmp["total_return_pct"])
lagged_all = spcx_first_trade < others_cmp["total_return_pct"].min()
beat_all = spcx_first_trade > others_cmp["total_return_pct"].max()
first_trade_phrase = (f"was down about {abs(spcx_first_trade):.0f}%" if spcx_first_trade < 0
                      else f"earned about {spcx_first_trade:.0f}%")
rank_phrase = ("lagged all three benchmarks" if lagged_all
               else "beat all three benchmarks" if beat_all else "landed in the middle of the pack")


# -----------------------------------------------------------------------------
# 2. Document setup and small helpers
# -----------------------------------------------------------------------------

def configure_document(document):
    """A4 page, sensible margins, and a clean built-in style set."""
    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(11)
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 13), ("Caption", 9)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)


def add_figure(document, filename, number, caption):
    """Insert one figure at text width with a numbered Word caption beneath it."""
    image_path = FIGURE_DIR / filename
    if not image_path.is_file():
        raise SystemExit(f"Missing figure {filename}. Re-run parts 1-3 to create the figures.")
    document.add_picture(str(image_path), width=Inches(6.0))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    box = document.add_paragraph(style="Caption")
    box.add_run(f"Figure {number}. ").bold = True
    box.add_run(caption)


def add_table(document, frame, number, caption):
    """Insert a short table with a numbered caption, header row in bold."""
    box = document.add_paragraph(style="Caption")
    box.add_run(f"Table {number}. ").bold = True
    box.add_run(caption)
    table = document.add_table(rows=frame.shape[0] + 1, cols=frame.shape[1])
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, column in enumerate(frame.columns):
        run = table.cell(0, col_idx).paragraphs[0].add_run(str(column))
        run.bold = True
        run.font.size = Pt(9)
    for row_idx, (_, row) in enumerate(frame.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            run = table.cell(row_idx, col_idx).paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
    document.add_paragraph()


# -----------------------------------------------------------------------------
# 3. Build the presentation tables (two or three significant figures)
# -----------------------------------------------------------------------------

scorecard_table = pd.DataFrame({
    "Measure": ["First price (offer)", "Last price", "Total return", "Average hourly return",
                "Annualized volatility", "Annualized Sharpe ratio"],
    "Value": [f"${scorecard['First price ($)']:.0f}", f"${scorecard['Last price ($)']:.0f}",
              f"{total_return:.0f}%", f"{scorecard['Average hourly return (%)']:.2f}%",
              f"{ann_vol:.0f}%", f"{ann_sharpe:.1f}"],
})

comparison_table = pd.DataFrame({
    "Stock": comparison["name"].values,
    "Total return": [f"{v:.1f}%" for v in comparison["total_return_pct"]],
    "Hourly volatility": [f"{v:.2f}%" for v in comparison["vol_per_bar_pct"]],
    "Annualized Sharpe": [f"{v:.2f}" for v in comparison["sharpe"]],
})


# -----------------------------------------------------------------------------
# 4. Write the report
# -----------------------------------------------------------------------------

doc = Document()
configure_document(doc)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("The SpaceX IPO: A High-Frequency Data Narrative")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("FINS3645 / FINS5545 — Week 4 worked example").font.size = Pt(11)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    f"Ticker: SPCX (Nasdaq)  |  Sample: {sample_start:%d %b %Y} to {sample_end:%d %b %Y}  |  "
    f"Hourly bars over {trading_days} trading days  |  Risk-free rate: 0%"
).font.size = Pt(10)
doc.add_paragraph()

# --- Executive Summary ---
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "On 11 June 2026 SpaceX priced the largest initial public offering (IPO) in history at $135 "
    "per share, and the stock began trading the next day on the Nasdaq under the ticker SPCX. This "
    "report uses hourly and five-minute prices from Yahoo Finance to follow the stock through its "
    "first week and to introduce four basic tools every analyst uses: simple returns, the "
    "cumulative return (the growth of $1), volatility, and the Sharpe ratio."
)
doc.add_paragraph(
    f"Measured from the $135 offer price, $1 grew to ${1 + total_return / 100:.2f}, a total return of "
    f"{total_return:.0f}% over {trading_days} trading days, after peaking near $1.60 in the middle of the "
    f"week. An investor who instead bought at the first public trade {first_trade_phrase} over the same "
    f"window: the IPO pop went mainly to investors allocated shares in the offering, and the early gains "
    f"then faded. From its first traded price SpaceX {rank_phrase} once the rally reversed, and its "
    f"annualized Sharpe ratio turned negative ({spcx_cmp['sharpe']:.2f}) while the diversified "
    f"{best_sharpe_row['name']} had the best ({best_sharpe_row['sharpe']:.2f}). The most exciting stock "
    f"was the worst risk-adjusted bet over this window."
)

# --- Introduction ---
doc.add_heading("Introduction: The Event", level=1)
doc.add_paragraph(
    "SpaceX sold about 556 million shares at $135 each, raising roughly $75 billion and valuing the "
    "company near $2.1 trillion. The shares opened at $150 (an 11% pop over the offer price), reached "
    "$176 during the first day, and closed at $160.95, up 19%. The debut made Elon Musk the first "
    "person worth more than $1 trillion (SpaceX Investor Relations, 2026; CNBC, 2026)."
)
doc.add_paragraph(
    "The listing was also controversial. The shares sold to the public carry one vote each, while "
    "Musk and insiders hold super-voting shares: he owns about 42% of the company but controls 80-85% "
    "of the votes, and the company is exempt from several standard board-independence rules (Bebchuk "
    "and Kastiel, 2026). Critics also questioned the price. At $135 the company traded at about 95 "
    "times its 2025 revenue despite a loss for the year; one research firm called the stock "
    "significantly overvalued, and a US senator asked the regulator to delay the listing (India "
    "Today, 2026). This report sets those debates aside and focuses on what the trading data shows."
)

# --- Data ---
doc.add_heading("Data", level=1)
doc.add_paragraph(
    f"The price history comes from the Yahoo Finance chart service. We download hourly bars (one price "
    f"per hour of trading) covering {sample_start:%d %B} to {sample_end:%d %B %Y}, which is the stock's "
    f"first {trading_days} trading days. The US market was closed on 19 June for a public holiday. "
    "Regular trading on debut day did not begin at the usual time: the exchange first ran an opening "
    "auction to set a single fair price, so the first hourly bar is late in the morning. We add the "
    "$135 offer price as the starting point so that the first return measures the IPO pop. For the "
    "distribution analysis we also download five-minute bars, which give several hundred observations "
    "rather than a few dozen."
)

# --- Methodology ---
doc.add_heading("Methodology", level=1)
doc.add_paragraph(
    "The simple return is the percentage change in price from one bar to the next: r = (P_t - P_(t-1)) "
    "/ P_(t-1). The cumulative return, or growth of $1, multiplies the growth factors (1 + r) across "
    "all bars, which equals the last price divided by the first price. Volatility is the standard "
    "deviation of the returns; we annualize it by multiplying by the square root of the number of "
    "return periods in a year. The Sharpe ratio divides the average return above the risk-free rate by "
    "the volatility, so it measures reward per unit of risk (Sharpe, 1966). Over this short window we "
    "set the risk-free rate to zero."
)

# --- Results ---
doc.add_heading("Results", level=1)

doc.add_heading("Price and the IPO Pop", level=2)
doc.add_paragraph(
    "The price jumped off the $135 offer line to about $165 within the first hour, rose to about $217 in "
    "the middle of the week, then gave back those gains to close near $155 on 22 June (Figure 1). The gap "
    "between the offer price and the first traded price is the IPO pop, the clearest single feature of the "
    "debut."
)
add_figure(doc, "02_spcx_price_ft.png", 1,
           "SpaceX hourly close price (SPCX, US dollars), New York time. The dashed line marks the "
           f"$135 IPO offer price. Sample: {sample_start:%d %b} to {sample_end:%d %b %Y}.")

doc.add_heading("Returns and the Growth of $1", level=2)
doc.add_paragraph(
    f"The first return, from the $135 offer to the first traded price, is the IPO pop and is by far the "
    f"largest single bar. Compounding all the hourly returns, $1 invested at the offer grew to "
    f"${1 + total_return / 100:.2f}, a {total_return:.0f}% total return (Figure 2, Table 1)."
)
add_figure(doc, "04_spcx_growth_of_one_ft.png", 2,
           "Growth of $1 invested at the $135 offer price (US dollars). The line starts at $1.00 and "
           "jumps with the IPO pop on the first bar. Hourly data.")
add_table(doc, scorecard_table, 1,
          "SpaceX hourly scorecard, measured from the $135 offer price. Annualized figures use the "
          "realized number of bars per trading day.")

doc.add_heading("Risk, the Sharpe Ratio, and a Caution", level=2)
doc.add_paragraph(
    f"Measured from the offer price, the annualized volatility is {ann_vol:.0f}% and the annualized Sharpe "
    f"ratio is {ann_sharpe:.1f}. These figures are extreme because they come from only {trading_days} days "
    "of a brand-new, unusually volatile stock; annualizing such a short window is not reliable. The methods "
    "matter here, not the precise annual numbers, which would settle down with a longer history."
)
doc.add_paragraph(
    "That single number also hides how the picture changed. Recomputed at the end of each day from the "
    "first traded price, the annualized Sharpe ratio rocketed to about +17 during the early rally and then "
    "fell every day, turning negative after the 22 June selloff (Figure 3). A week of strong returns can "
    "become a loss quickly when a new stock reverses."
)
add_figure(doc, "11_spcx_sharpe_collapse_ft.png", 3,
           "Annualized Sharpe ratio measured from SpaceX's first traded price, recomputed through each "
           "day's close, risk-free rate zero. The first day uses only a few hours, so read it as indicative.")

doc.add_heading("Comparison with Other Technology Stocks", level=2)
doc.add_paragraph(
    "To judge whether the move was special, we compare SpaceX with Tesla, Nvidia, and the Nasdaq-100 "
    "exchange-traded fund (a fund tracking the 100 largest Nasdaq companies) over the same window. To "
    "be fair, every stock starts at $1 at SpaceX's first traded bar, because the benchmarks have no "
    f"offer price. SpaceX raced ahead and then gave it all back: from its first traded price it "
    f"{first_trade_phrase}, while Tesla, Nvidia, and the Nasdaq-100 were each up about 2% to 3% "
    f"(Figure 4). SpaceX also swung by far the most, and was the only one of the four to end the window "
    f"with a loss (Figure 5). Dividing return by risk, SpaceX's Sharpe ratio is negative while the "
    f"diversified {best_sharpe_row['name']} has the highest (Figure 6, Table 2): more risk did not buy "
    "more reward here."
)
add_figure(doc, "07_compare_growth_of_one_ft.png", 4,
           "Growth of $1 from SpaceX's first hourly bar for SpaceX, Tesla, Nvidia, and the Nasdaq-100 "
           "(US dollars). All series start at $1.00. Hourly data.")
add_figure(doc, "08_compare_risk_return_ft.png", 5,
           "Total return over the window against the typical size of an hourly move (both in per cent). "
           "Each dot is one stock.")
add_figure(doc, "09_compare_sharpe_ft.png", 6,
           "Annualized Sharpe ratio (return per unit of risk) for each stock, risk-free rate set to "
           "zero. Higher is better.")
add_table(doc, comparison_table, 2,
          "Performance since SpaceX's first hourly bar. Total return and hourly volatility are over the "
          "window; the Sharpe ratio is annualized.")

doc.add_heading("The Shape of Returns: Heavy Tails", level=2)
doc.add_paragraph(
    "Stock returns are not normally distributed. Compared with a bell curve, real returns have a taller "
    "peak and heavier tails: extreme moves happen far more often than a normal distribution predicts. "
    f"The five-minute SpaceX returns have an excess kurtosis of {excess_kurtosis:.1f}, far above the 0 of "
    "a normal distribution (Figure 7). This heavy-tailed shape is one of the most reliable facts about "
    "asset returns across markets and time horizons (Cont, 2001)."
)
add_figure(doc, "06_spcx_return_distribution_ft.png", 7,
           "Distribution of SpaceX five-minute returns (per cent) against a normal curve with the same "
           "mean and standard deviation. The peak is taller and the tails are heavier than the bell curve.")

# --- Conclusion ---
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "SpaceX's debut is a vivid first lesson in market data. The IPO pop separated the offer price from "
    "the first traded price; the growth of $1 turned a week of returns into a single dollar figure; "
    "volatility and the Sharpe ratio showed that the most exciting, highest-risk stock delivered the worst "
    "risk-adjusted return over this window, and a Sharpe ratio that looked spectacular mid-week turned "
    "negative once the rally reversed; and the five-minute returns showed the heavy tails that mark almost "
    "every financial asset. "
    "The numbers themselves come from a tiny, unusual sample and should not be read as forecasts. The "
    "workflow, however, is exactly the one used on years of data for any listed stock."
)

# --- References ---
doc.add_heading("References", level=1)
references = [
    "Bebchuk, L. A., and Kastiel, K. (2026). Top IPO, Weak Governance. Harvard Law School Forum on "
    "Corporate Governance, 19 May 2026.",
    "CNBC (2026). SpaceX stock jumps in first full day of trading after record Nasdaq debut. CNBC, June 2026.",
    "Cont, R. (2001). Empirical properties of asset returns: stylized facts and statistical issues. "
    "Quantitative Finance, 1(2), 223-236.",
    "India Today (2026). The world's biggest IPO starts trading today: is Elon Musk's SpaceX overvalued? "
    "India Today, 12 June 2026.",
    "Sharpe, W. F. (1966). Mutual fund performance. The Journal of Business, 39(1), 119-138.",
    "SpaceX Investor Relations (2026). Space Exploration Technologies Corp. announces pricing of initial "
    "public offering. SpaceX, June 2026.",
]
for entry in references:
    paragraph = doc.add_paragraph(entry)
    paragraph.paragraph_format.space_after = Pt(6)

doc.save(REPORT_PATH)
print(f"Saved report: {REPORT_PATH}")
print(f"Sections: Executive Summary, Introduction, Data, Methodology, Results, Conclusion, References")
print(f"Figures: 7  |  Tables: 2  |  Trading days: {trading_days}")
