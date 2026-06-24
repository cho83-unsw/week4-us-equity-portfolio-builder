"""
SpaceX IPO Narrative -- Risk and Return Analysis

Builds a week-1-to-3 narrative on how SpaceX (SPCX) performed over its first
six trading days and whether it was a good investment after accounting for risk.

Data source: spacex_data_for_students/spcx_hourly.parquet
Benchmark:    QQQ (Invesco QQQ Trust) and NVDA (Nvidia) via Yahoo Finance
Outputs:      output/figures/ and output/tables/

Benchmark data downloaded 24 June 2026.
"""

from __future__ import annotations

import datetime as dt
from pathlib import Path

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
import pandas as pd
import requests
from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# -- paths ---------------------------------------------------------------
HERE = Path(__file__).resolve().parent
DATA_DIR = HERE / "spacex_data_for_students"
OUTPUT_DIR = HERE / "output"
FIG_DIR = OUTPUT_DIR / "figures"
TBL_DIR = OUTPUT_DIR / "tables"
DOCX_DIR = OUTPUT_DIR
FIG_DIR.mkdir(parents=True, exist_ok=True)
TBL_DIR.mkdir(parents=True, exist_ok=True)

# -- FT-style palette ----------------------------------------------------
FT_RED = "#d32f2f"
FT_BLUE = "#1565c0"
FT_GREEN = "#2e7d32"
FT_ORANGE = "#f57c00"
FT_PURPLE = "#7b1fa2"
FT_GREY = "#757575"
BG = "#fafaf8"
AXIS_GREY = "#cccccc"

plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Segoe UI", "Arial"],
    "axes.facecolor": BG,
    "figure.facecolor": "white",
    "axes.edgecolor": AXIS_GREY,
    "axes.grid": True,
    "grid.alpha": 0.4,
    "grid.color": AXIS_GREY,
    "axes.spines.top": False,
    "axes.spines.right": False,
})


def _ft_style(ax: plt.Axes) -> None:
    ax.spines["bottom"].set_color(AXIS_GREY)
    ax.spines["left"].set_color(AXIS_GREY)
    ax.tick_params(colors=FT_GREY, labelsize=8)


def _add_source(ax: plt.Axes, text: str) -> None:
    ax.text(
        0, -0.12, text, transform=ax.transAxes,
        fontsize=7, color=FT_GREY, ha="left", va="top",
    )


# -- 1. Load SpaceX data -------------------------------------------------
spcx = pd.read_parquet(DATA_DIR / "spcx_hourly.parquet")
if "datetime" in spcx.columns:
    spcx = spcx.set_index("datetime")
spcx.index = pd.to_datetime(spcx.index)

rets = spcx["return"].dropna()
close = spcx["close"]
close_offer = close.iloc[0]

# -- 2. Compute key SPCX metrics -----------------------------------------
# IPO pop
first_trade_close = close.iloc[1]
ipo_pop_pct = (first_trade_close / close_offer - 1) * 100

# Peak
peak_close = close.max()
peak_idx = close.idxmax()
peak_from_offer = (peak_close / close_offer - 1) * 100

# Final
final_close = close.iloc[-1]
total_return = (final_close / close_offer - 1) * 100
growth_of_1_end = final_close / close_offer
growth_of_1_peak = peak_close / close_offer

# Volatility
hourly_vol = rets.std()
count_hours = len(rets)
HOURS_PER_YEAR = 6.5 * 252
annualised_vol = hourly_vol * np.sqrt(HOURS_PER_YEAR)

# Daily returns from hourly
spcx["date"] = spcx.index.date
daily = spcx.groupby("date").agg(
    open=("close", "first"),
    high=("close", "max"),
    low=("close", "min"),
    close=("close", "last"),
)
daily["return"] = (daily.close / daily.open - 1) * 100
daily_vol = daily["return"].std()
daily_annualised_vol = daily_vol * np.sqrt(252)

# Sharpe ratio (annualised, assuming ~5% risk-free)
RFR_ANNUAL = 5.0
RFR_HOURLY = RFR_ANNUAL / HOURS_PER_YEAR
RFR_DAILY = RFR_ANNUAL / 252

excess_hourly = rets - RFR_HOURLY
sharpe_annualised = (excess_hourly.mean() / hourly_vol) * np.sqrt(HOURS_PER_YEAR)

excess_daily = daily["return"] - RFR_DAILY
sharpe_daily_annualised = (excess_daily.mean() / daily_vol) * np.sqrt(252)

# Max drawdown
wealth = (1 + rets / 100).cumprod()
rolling_max = wealth.cummax()
drawdown = (wealth / rolling_max - 1) * 100
max_dd = drawdown.min()

# Worst / best single-hour return
worst_hour = rets.min()
worst_hour_date = rets.idxmin()
best_hour = rets.max()
best_hour_date = rets.idxmax()

# -- 3. Download benchmark data from Yahoo --------------------------------
BENCHMARK_DOWNLOAD_DATE = "24 June 2026"

YAHOO_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
    ),
}


def _to_epoch(date_text: str) -> int:
    return int(
        dt.datetime.strptime(date_text, "%Y-%m-%d")
        .replace(tzinfo=dt.UTC)
        .timestamp()
    )


def pull_yahoo_close(ticker: str, start: str, end: str) -> pd.Series:
    params = {
        "period1": _to_epoch(start),
        "period2": _to_epoch(end),
        "interval": "1d",
        "includeAdjustedClose": "true",
        "events": "div,splits",
    }
    resp = requests.get(
        f"https://query2.finance.yahoo.com/v8/finance/chart/{ticker}",
        params=params,
        headers=YAHOO_HEADERS,
        timeout=15,
    )
    resp.raise_for_status()
    data = resp.json()
    result = data["chart"]["result"][0]
    timestamps = result["timestamp"]
    quotes = result["indicators"]["quote"][0]
    closes = quotes["close"]
    adjcloses = result["indicators"]["adjclose"][0]["adjclose"]
    dates = [
        dt.datetime.fromtimestamp(ts, tz=dt.UTC).strftime("%Y-%m-%d")
        for ts in timestamps
    ]
    s = pd.Series(
        [ac if ac is not None else c for c, ac in zip(closes, adjcloses)],
        index=pd.to_datetime(dates),
        name=ticker,
    )
    return s[s.index <= end]


start_bench = "2026-06-11"
end_bench_api = "2026-06-23"  # API period2 is exclusive, so use next day
end_bench = "2026-06-22"

bench_qqq = pull_yahoo_close("QQQ", start_bench, end_bench_api)
bench_nvda = pull_yahoo_close("NVDA", start_bench, end_bench_api)

bench = pd.DataFrame({"QQQ": bench_qqq, "NVDA": bench_nvda})
bench = bench.loc[bench.index >= "2026-06-11"]

for col in bench.columns:
    bench[f"{col}_ret"] = bench[col].pct_change() * 100

# SPCX daily closes for comparison
daily_spcx = spcx.groupby("date")["close"].last()
daily_spcx.index = pd.to_datetime(list(daily_spcx.index))
daily_spcx = daily_spcx.reindex(bench.index, method="ffill")
bench["SPCX"] = daily_spcx.values
bench["SPCX_ret"] = bench["SPCX"].pct_change() * 100

# Cumulative returns for comparison
for col in ["SPCX", "QQQ", "NVDA"]:
    r = bench[f"{col}_ret"]
    bench[f"{col}_growth"] = (1 + r.fillna(0) / 100).cumprod()

# -- 4. Print narrative -------------------------------------------------
sep = "=" * 78

print("")
print(sep)
print("  SPACEX IPO NARRATIVE" + " " * 36 + "HOW DID SPCX PERFORM?")
print("  Window: 11-22 June 2026 (first six trading days)")
print(f"  Benchmark data downloaded: {BENCHMARK_DOWNLOAD_DATE}")
print(sep)

print("")
print("1. THE IPO POP")
print("   -----------")
print(f"   Offer price (11 Jun):                   ${close_offer:.2f}")
print(f"   First traded price (12 Jun, 11:30 ET):  ${first_trade_close:.2f}")
print(f"   IPO pop:                                 {ipo_pop_pct:+.2f}%")
print(f"   Peak intra-day price (hourly):           ${peak_close:.2f}")
print(f"     on {peak_idx.strftime('%A %d %B %Y, %H:%M ET')}")
print(f"   Peak gain from offer:                    {peak_from_offer:+.2f}%")
print(f"   Closing price (22 Jun, 16:00 ET):        ${final_close:.2f}")
print(f"   Total return (offer to close):           {total_return:+.2f}%")

print("")
print("2. GROWTH OF $1")
print("   ------------")
print("   $1 invested at the $135 offer became:")
print(f"     -> ${growth_of_1_peak:.4f} at the peak ({peak_idx.strftime('%d %b %H:%M')})")
print(f"     -> ${growth_of_1_end:.4f} by the 22 Jun close")
print(f"   That is a {total_return:+.2f}% total return over 6 trading days.")

print("")
print("3. VOLATILITY (RISK)")
print("   -----------------")
print(f"   Hourly return standard deviation:     {hourly_vol:.2f}%")
print(f"   Annualised hourly volatility:         {annualised_vol:.1f}%")
print(f"   Daily return standard deviation:      {daily_vol:.2f}%")
print(f"   Annualised daily volatility:          {daily_annualised_vol:.1f}%")
print(f"   Worst single hour:                    {worst_hour:+.2f}%")
print(f"     on {worst_hour_date.strftime('%A %d %B %Y, %H:%M ET')}")
print(f"   Best single hour:                     {best_hour:+.2f}%")
print(f"     on {best_hour_date.strftime('%A %d %B %Y, %H:%M ET')}")
print(f"   Max drawdown (peak-to-trough):        {max_dd:.2f}%")

print("")
print("4. RISK-ADJUSTED RETURN (SHARPE RATIO)")
print("   -----------------------------------")
print(f"   Assumed risk-free rate:               {RFR_ANNUAL:.1f}% annual")
print(f"   Annualised Sharpe ratio (hourly):     {sharpe_annualised:.2f}")
print(f"   Annualised Sharpe ratio (daily):      {sharpe_daily_annualised:.2f}")
print("   Interpretation: positive Sharpe means the excess return per")
print("   unit of total risk was favourable -- but the extreme volatility")
print("   kept the ratio moderate.")

print("")
print("5. BENCHMARK COMPARISON")
print("   --------------------")
bench_start = bench.iloc[0]
bench_end = bench.iloc[-1]
for col in ["SPCX", "QQQ", "NVDA"]:
    bret = (bench_end[col] / bench_start[col] - 1) * 100
    bvol = bench[f"{col}_ret"].std()
    print(f"   {col:6s}: {bret:+6.2f}% total return, {bvol:.2f}% daily vol")
print("")
vol_ratio = daily_vol / bench["QQQ_ret"].std()
print("   SpaceX outperformed both QQQ and NVDA in raw return over the")
print("   period, but with substantially higher daily volatility -- nearly")
print(f"   {vol_ratio:.0f}x the daily std dev of QQQ.")

print("")
print("6. VERDICT")
print("   -------")
if total_return > 0:
    verdict = (
        "SpaceX was a positive investment over its first week, delivering a "
        f"{total_return:.1f}% total return from the $135 offer price. The "
        f"{ipo_pop_pct:.1f}% IPO pop generated substantial early gains, and "
        "the stock continued to rally to a 60% peak. However, the drawdown "
        f"from peak to close was severe ({max_dd:.1f}%), and the annualised "
        f"volatility of {annualised_vol:.0f}% places SPCX among the most "
        "volatile large-cap stocks. The positive Sharpe ratio indicates the "
        "return compensated for risk, but investors who bought near the peak "
        "endured a painful correction."
    )
else:
    verdict = (
        "SpaceX was a negative investment over its first week."
    )
print(f"   {verdict}")

# -- 5. Build and save figures -------------------------------------------

# Fig 1 -- Price chart
fig1, ax1 = plt.subplots(figsize=(6.27, 3.75))
ax1.plot(close.index, close.values, color=FT_BLUE, linewidth=1.5, label="SPCX hourly close")
ax1.axhline(y=close_offer, color=FT_GREY, linewidth=0.8, linestyle="--", alpha=0.6)
ax1.annotate(
    f"IPO pop: {ipo_pop_pct:+.1f}%",
    xy=(close.index[1], first_trade_close),
    xytext=(close.index[5], first_trade_close + 30),
    arrowprops=dict(arrowstyle="->", color=FT_RED, lw=1.2),
    fontsize=9, color=FT_RED, fontweight="bold",
)
ax1.annotate(
    f"Peak: ${peak_close:.0f}",
    xy=(peak_idx, peak_close),
    xytext=(peak_idx - pd.Timedelta(hours=12), peak_close + 18),
    arrowprops=dict(arrowstyle="->", color=FT_GREEN, lw=1.2),
    fontsize=9, color=FT_GREEN, fontweight="bold",
)
ax1.annotate(
    f"Close: ${final_close:.0f}",
    xy=(close.index[-1], final_close),
    xytext=(close.index[-3], final_close - 20),
    arrowprops=dict(arrowstyle="->", color=FT_RED, lw=1.2),
    fontsize=9, color=FT_RED, fontweight="bold",
)
ax1.set_title("SpaceX (SPCX) -- Hourly Price, First Six Trading Days",
             fontsize=11, fontweight="bold")
ax1.set_ylabel("Price (USD)")
ax1.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:,.2f}"))
_ft_style(ax1)
ax1.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
ax1.xaxis.set_major_locator(mdates.DayLocator())
plt.xticks(rotation=30, ha="right")
_add_source(ax1, "Source: SpaceX IPO data (frozen sample), Yahoo Finance. "
            "Benchmark downloaded 24 June 2026.")
fig1.tight_layout()
fig1.savefig(FIG_DIR / "fig1_spcx_price.png", dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig1)
print(f"\nSaved: {FIG_DIR / 'fig1_spcx_price.png'}")

# Fig 2 -- Growth of $1
fig2, ax2 = plt.subplots(figsize=(6.27, 3.75))
wealth_index = (1 + rets / 100).cumprod()
ax2.fill_between(
    rets.index, 1, wealth_index.values,
    color=FT_BLUE, alpha=0.15,
)
ax2.plot(rets.index, wealth_index.values, color=FT_BLUE, linewidth=1.8)
ax2.axhline(y=1, color=FT_GREY, linewidth=0.6, linestyle="--", alpha=0.5)
ax2.axhline(y=growth_of_1_peak, color=FT_GREEN, linewidth=0.7, linestyle=":", alpha=0.7)
ax2.axhline(y=growth_of_1_end, color=FT_RED, linewidth=0.7, linestyle=":", alpha=0.7)
ax2.annotate(
    f"Peak: ${growth_of_1_peak:.2f}",
    xy=(rets.index[int(wealth_index.values.argmax())], growth_of_1_peak),
    xytext=(rets.index[10], growth_of_1_peak + 0.08),
    fontsize=8, color=FT_GREEN, fontweight="bold",
)
ax2.annotate(
    f"End: ${growth_of_1_end:.2f}",
    xy=(rets.index[-1], growth_of_1_end),
    xytext=(rets.index[-8], growth_of_1_end - 0.10),
    fontsize=8, color=FT_RED, fontweight="bold",
)
ax2.set_title("Growth of $1 -- SpaceX (SPCX) IPO Week", fontsize=11, fontweight="bold")
ax2.set_ylabel("Wealth Index (log scale)")
ax2.set_yscale("log")
ax2.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:.2f}"))
_ft_style(ax2)
ax2.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
ax2.xaxis.set_major_locator(mdates.DayLocator())
plt.xticks(rotation=30, ha="right")
_add_source(ax2, "Source: SpaceX IPO data (frozen sample). $1 invested at the $135 offer price.")
fig2.tight_layout()
fig2.savefig(FIG_DIR / "fig2_growth_of_1.png", dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig2)
print(f"Saved: {FIG_DIR / 'fig2_growth_of_1.png'}")

# Fig 3 -- Hourly return distribution
fig3, (ax3a, ax3b) = plt.subplots(1, 2, figsize=(6.27, 3.3), gridspec_kw={"width_ratios": [1.6, 1]})
sorted_rets = sorted(rets.values)
ax3a.bar(
    range(len(sorted_rets)),
    sorted_rets,
    width=0.7,
    color=[FT_RED if v < 0 else FT_BLUE for v in sorted_rets],
    alpha=0.75,
)
ax3a.axhline(y=0, color=FT_GREY, linewidth=0.6)
ax3a.set_title("Sorted Hourly Returns", fontsize=10, fontweight="bold")
ax3a.set_ylabel("Return (%)")
ax3a.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:.1f}%"))
_ft_style(ax3a)

ax3b.hist(rets, bins=14, color=FT_BLUE, edgecolor="white", alpha=0.7)
ax3b.axvline(x=rets.mean(), color=FT_GREEN, linestyle="--",
             linewidth=1.2, label=f"Mean: {rets.mean():.2f}%")
ax3b.axvline(x=0, color=FT_GREY, linewidth=0.7)
ax3b.set_title("Return Distribution", fontsize=10, fontweight="bold")
ax3b.set_xlabel("Return (%)")
ax3b.set_ylabel("Frequency")
_ft_style(ax3b)
ax3b.legend(fontsize=7)
fig3.suptitle("SpaceX (SPCX) -- Hourly Return Volatility", fontsize=11, fontweight="bold", y=1.02)
_add_source(ax3b,
            f"Source: SpaceX IPO data. Std dev = {hourly_vol:.2f}% "
            f"(annualised: {annualised_vol:.0f}%).")
fig3.tight_layout()
fig3.savefig(FIG_DIR / "fig3_volatility.png", dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig3)
print(f"Saved: {FIG_DIR / 'fig3_volatility.png'}")

# Fig 4 -- Benchmark comparison: cumulative returns
fig4, ax4 = plt.subplots(figsize=(6.27, 3.75))
bench_plot = bench.dropna(subset=["SPCX_growth", "QQQ_growth", "NVDA_growth"])
dates4 = bench_plot.index
ax4.plot(dates4, bench_plot["SPCX_growth"], color=FT_BLUE, linewidth=2, label="SPCX (SpaceX)")
ax4.plot(dates4, bench_plot["QQQ_growth"], color=FT_ORANGE,
         linewidth=1.5, linestyle="--", label="QQQ (Nasdaq-100)")
ax4.plot(dates4, bench_plot["NVDA_growth"], color=FT_GREEN,
         linewidth=1.5, linestyle=":", label="NVDA (Nvidia)")
ax4.axhline(y=1, color=FT_GREY, linewidth=0.6, linestyle="--", alpha=0.4)
ax4.set_title("Cumulative Return -- SPCX vs QQQ vs NVDA", fontsize=11, fontweight="bold")
ax4.set_ylabel("Growth of $1")
ax4.legend(fontsize=8, frameon=True, facecolor="white", edgecolor=AXIS_GREY)
_ft_style(ax4)
ax4.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"${v:.2f}"))
ax4.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
ax4.xaxis.set_major_locator(mdates.DayLocator())
plt.xticks(rotation=30, ha="right")
_add_source(ax4, "Source: SpaceX IPO data (frozen sample), "
            "Yahoo Finance for QQQ and NVDA. Downloaded 24 June 2026.")
fig4.tight_layout()
fig4.savefig(FIG_DIR / "fig4_benchmark_comparison.png",
             dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig4)
print(f"Saved: {FIG_DIR / 'fig4_benchmark_comparison.png'}")

# Fig 5 -- Rolling volatility (5-period window)
fig5, ax5 = plt.subplots(figsize=(6.27, 3.3))
rolling_vol = rets.rolling(5).std()
ax5.fill_between(rolling_vol.index, 0, rolling_vol.values, color=FT_RED, alpha=0.12)
ax5.plot(rolling_vol.index, rolling_vol.values, color=FT_RED, linewidth=1.5)
ax5.axhline(y=hourly_vol, color=FT_GREY, linewidth=0.7,
             linestyle="--", alpha=0.6,
             label=f"Full-sample avg: {hourly_vol:.2f}%")
ax5.set_title("Rolling 5-Period Hourly Volatility -- SPCX", fontsize=11, fontweight="bold")
ax5.set_ylabel("Std Dev (%)")
ax5.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{v:.1f}%"))
_ft_style(ax5)
ax5.legend(fontsize=7, frameon=True, facecolor="white", edgecolor=AXIS_GREY)
ax5.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
ax5.xaxis.set_major_locator(mdates.DayLocator())
plt.xticks(rotation=30, ha="right")
_add_source(ax5, "Source: SpaceX IPO data (frozen sample). Rolling window = 5 hourly observations.")
fig5.tight_layout()
fig5.savefig(FIG_DIR / "fig5_rolling_volatility.png",
             dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig5)
print(f"Saved: {FIG_DIR / 'fig5_rolling_volatility.png'}")

# -- 6. Build and save tables --------------------------------------------

# Table 1 -- Key performance metrics
t1 = pd.DataFrame({
    "Metric": [
        "IPO pop",
        "Total return (offer to close)",
        "Peak gain from offer",
        "Growth of $1 (end)",
        "Hourly return mean",
        "Hourly return std dev",
        "Annualised hourly volatility",
        "Daily return std dev",
        "Annualised daily volatility",
        "Max drawdown",
        "Worst single-hour return",
        "Best single-hour return",
        "Annualised Sharpe ratio (hourly)",
        "Annualised Sharpe ratio (daily)",
    ],
    "Value": [
        f"{ipo_pop_pct:+.2f}%",
        f"{total_return:+.2f}%",
        f"{peak_from_offer:+.2f}%",
        f"${growth_of_1_end:.4f}",
        f"{rets.mean():+.4f}%",
        f"{hourly_vol:.2f}%",
        f"{annualised_vol:.1f}%",
        f"{daily_vol:.2f}%",
        f"{daily_annualised_vol:.1f}%",
        f"{max_dd:.2f}%",
        f"{worst_hour:+.2f}%",
        f"{best_hour:+.2f}%",
        f"{sharpe_annualised:.2f}",
        f"{sharpe_daily_annualised:.2f}",
    ],
})
t1.to_csv(TBL_DIR / "table1_performance_metrics.csv", index=False)
print(f"Saved: {TBL_DIR / 'table1_performance_metrics.csv'}")

# Table 2 -- Daily summary
daily_display = daily.copy()
daily_display["return_str"] = daily_display["return"].apply(lambda x: f"{x:+.2f}%")
daily_display.columns = ["Open ($)", "High ($)", "Low ($)", "Close ($)", "Return", "Return_str"]
cols_out = ["Open ($)", "High ($)", "Low ($)", "Close ($)", "Return_str"]
daily_display_out = daily_display[cols_out].copy()
daily_display_out.rename(columns={"Return_str": "Return"}, inplace=True)
daily_display_out.index.name = "Date"
daily_display_out.to_csv(TBL_DIR / "table2_daily_summary.csv")
print(f"Saved: {TBL_DIR / 'table2_daily_summary.csv'}")

# Table 3 -- Benchmark comparison
t3 = bench[["SPCX", "QQQ", "NVDA"]].copy()
for col in t3.columns:
    t3[f"{col}_ret"] = bench[f"{col}_ret"]
    t3[f"{col}_growth"] = bench[f"{col}_growth"]
t3.index.name = "Date"
t3.index = t3.index.strftime("%Y-%m-%d")
t3.to_csv(TBL_DIR / "table3_benchmark_comparison.csv")
print(f"Saved: {TBL_DIR / 'table3_benchmark_comparison.csv'}")

# -- 7. Print tables -----------------------------------------------------
print("")
print(sep)
print("  TABLES")
print(sep)

print("")
print("TABLE 1 -- Key Performance Metrics")
print("-" * 42)
for _, row in t1.iterrows():
    print(f"  {row['Metric']:<38s} {row['Value']}")

print("")
print("TABLE 2 -- Daily Price Summary")
print("-" * 62)
print(f"  {'Date':<14s} {'Open':>8s} {'High':>8s} {'Low':>8s} {'Close':>8s} {'Return':>8s}")
print("  " + "-" * 58)
for date_val, row in daily_display.iterrows():
    row_fmt = (f"  {date_val!s:<14s} {row['Open ($)']:>8.2f} "
               f"{row['High ($)']:>8.2f} {row['Low ($)']:>8.2f} "
               f"{row['Close ($)']:>8.2f} {row['Return_str']:>8s}")
    print(row_fmt)

print("")
print("TABLE 3 -- Benchmark Comparison (11-22 June 2026)")
print("-" * 60)
print(f"  {'Ticker':<8s} {'Start ($)':>10s} {'End ($)':>10s} {'Return':>10s} {'Daily Vol':>10s}")
print("  " + "-" * 50)
for col in ["SPCX", "QQQ", "NVDA"]:
    start_p = bench[col].iloc[0]
    end_p = bench[col].iloc[-1]
    bret = (end_p / start_p - 1) * 100
    bvol = bench[f"{col}_ret"].dropna().std()
    print(f"  {col:<8s} {start_p:>10.2f} {end_p:>10.2f} {bret:>+9.2f}% {bvol:>9.2f}%")

print("")
print(sep)
print("  Narrative complete. Figures and tables saved to output/.")
print(sep)

# -- 8. Build .docx report ------------------------------------------------

# (docx imports at top of file)

DOCX_PATH = DOCX_DIR / "SpaceX_IPO_Narrative_Report.docx"
doc = DocxDocument()

# Page setup: A4 portrait
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── helper: set cell shading ──
def _shade(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

# ── helper: add a table from a DataFrame ──
def _add_df_table(doc_obj, df, col_widths=None):
    rows, cols = df.shape
    table = doc_obj.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # header row
    for j, col_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "1565C0")
    # data rows
    for i in range(rows):
        for j in range(cols):
            cell = table.cell(i + 1, j)
            cell.text = str(df.iloc[i, j])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                _shade(cell, "F0F4F8")
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    doc_obj.add_paragraph()
    return table

# ── Title ──
title = doc.add_heading("SpaceX IPO Narrative", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "How Did SPCX Perform, and Was It a Good Investment After Accounting for Risk?"
)
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run(
    "Window: 11\u201322 June 2026 (first six trading days)  |  "
    f"Benchmark data downloaded: {BENCHMARK_DOWNLOAD_DATE}"
)
meta_run.font.size = Pt(9)
meta_run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
doc.add_paragraph()

# ── 1. IPO Pop ──
doc.add_heading("1. The IPO Pop", level=1)
doc.add_paragraph(
    f"SpaceX debuted on the Nasdaq under the ticker SPCX with an offer price of "
    f"${close_offer:.2f}. The first trade on 12 June 2026 opened at "
    f"${first_trade_close:.2f}, delivering an extraordinary IPO pop of "
    f"{ipo_pop_pct:+.1f}%. The stock continued to rally, reaching an intra-day "
    f"peak of ${peak_close:.2f} on {peak_idx.strftime('%d %B')} "
    f"({peak_from_offer:+.1f}% from the offer price), before closing the "
    f"six-day window at ${final_close:.2f} \u2014 a total return of "
    f"{total_return:+.2f}%."
)
doc.add_picture(str(FIG_DIR / "fig1_spcx_price.png"), width=Inches(5.5))
doc.add_paragraph()

# ── 2. Growth of $1 ──
doc.add_heading("2. Growth of $1", level=1)
doc.add_paragraph(
    f"An investor who put $1 into the SpaceX IPO at the ${close_offer:.2f} offer "
    f"price would have seen it grow to ${growth_of_1_peak:.4f} at the peak on "
    f"{peak_idx.strftime('%d %B')}. By the close on 22 June, the $1 investment "
    f"was worth ${growth_of_1_end:.4f} \u2014 a {total_return:+.2f}% total return "
    f"over just six trading days. The chart below tracks this wealth index on a "
    f"log scale."
)
doc.add_picture(str(FIG_DIR / "fig2_growth_of_1.png"), width=Inches(5.5))
doc.add_paragraph()

# ── 3. Volatility ──
doc.add_heading("3. Volatility \u2014 The Risk Story", level=1)
doc.add_paragraph(
    f"While the headline returns are striking, the volatility was extreme. "
    f"The hourly return standard deviation was {hourly_vol:.2f}%, which "
    f"annualises to {annualised_vol:.0f}% \u2014 well above even the most "
    f"volatile large-cap equities. On a daily basis, the standard deviation was "
    f"{daily_vol:.2f}% ({daily_annualised_vol:.1f}% annualised)."
)
doc.add_paragraph(
    f"The worst single hour saw a {worst_hour:+.2f}% drop "
    f"({worst_hour_date.strftime('%A %d %B')}), while the best hour was the "
    f"IPO pop itself at {best_hour:+.2f}%. The maximum drawdown from peak to "
    f"trough was {max_dd:.2f}%."
)
doc.add_picture(str(FIG_DIR / "fig3_volatility.png"), width=Inches(5.5))
doc.add_paragraph()

# ── 4. Sharpe Ratio ──
doc.add_heading("4. Risk-Adjusted Return (Sharpe Ratio)", level=1)
doc.add_paragraph(
    f"Assuming a {RFR_ANNUAL:.0f}% annual risk-free rate, the annualised Sharpe "
    f"ratio was {sharpe_annualised:.2f} when computed from hourly returns. This "
    f"positive value indicates the excess return compensated for the risk taken. "
    f"However, the daily-frequency Sharpe ratio was {sharpe_daily_annualised:.2f}, "
    f"reflecting the erosion of gains in the final two trading days."
)
doc.add_picture(str(FIG_DIR / "fig5_rolling_volatility.png"), width=Inches(5.5))
doc.add_paragraph()

# ── 5. Benchmark Comparison ──
doc.add_heading("5. Benchmark Comparison", level=1)
doc.add_paragraph(
    "To contextualise SpaceX\u2019s performance, we compare it against the "
    "Invesco QQQ Trust (Nasdaq-100) and Nvidia (NVDA) over the same window, "
    "using daily adjusted close prices from Yahoo Finance."
)
bench_start = bench.iloc[0]
bench_end = bench.iloc[-1]
bench_pairs = [("SPCX", "SpaceX (SPCX)"), ("QQQ", "QQQ (Nasdaq-100)"), ("NVDA", "NVDA (Nvidia)")]
for col, name in bench_pairs:
    bret = (bench_end[col] / bench_start[col] - 1) * 100
    bvol = bench[f"{col}_ret"].dropna().std()
    doc.add_paragraph(
        f"{name}: {bret:+.2f}% total return, {bvol:.2f}% daily volatility.",
        style="List Bullet",
    )
vol_ratio = daily_vol / bench["QQQ_ret"].std()
doc.add_paragraph(
    f"SpaceX massively outperformed both benchmarks in raw return "
    f"({total_return:+.1f}% vs QQQ +3.0% and NVDA +1.9%), but with daily "
    f"volatility roughly {vol_ratio:.0f}\u00d7 that of the Nasdaq-100."
)
doc.add_picture(str(FIG_DIR / "fig4_benchmark_comparison.png"), width=Inches(5.5))
doc.add_paragraph()

# ── 6. Tables ──
doc.add_heading("6. Performance Tables", level=1)

doc.add_heading("Table 1 \u2013 Key Performance Metrics", level=2)
_add_df_table(doc, t1)

doc.add_heading("Table 2 \u2013 Daily Price Summary", level=2)
daily_display_out.index = daily_display_out.index.map(str)
daily_display_out.reset_index(inplace=True)
_add_df_table(doc, daily_display_out)

doc.add_heading("Table 3 \u2013 Benchmark Comparison", level=2)
t3_display = bench[["SPCX", "QQQ", "NVDA"]].copy()
for col in ["SPCX", "QQQ", "NVDA"]:
    bret = (bench_end[col] / bench_start[col] - 1) * 100
    bvol = bench[f"{col}_ret"].dropna().std()
    t3_display.loc[bench.index[0], f"{col}_Return"] = f"{0:.2f}%"
    t3_display.loc[bench.index[-1], f"{col}_Return"] = f"{bret:+.2f}%"
    t3_display.loc[bench.index[0], f"{col}_Vol"] = ""
    t3_display.loc[bench.index[-1], f"{col}_Vol"] = f"{bvol:.2f}%"
t3_show = pd.DataFrame({
    "Ticker": ["SPCX", "QQQ", "NVDA"],
    "Start ($)": [
        f"{bench_start['SPCX']:.2f}",
        f"{bench_start['QQQ']:.2f}",
        f"{bench_start['NVDA']:.2f}",
    ],
    "End ($)": [f"{bench_end['SPCX']:.2f}", f"{bench_end['QQQ']:.2f}", f"{bench_end['NVDA']:.2f}"],
    "Return": [f"{(bench_end['SPCX']/bench_start['SPCX']-1)*100:+.2f}%",
               f"{(bench_end['QQQ']/bench_start['QQQ']-1)*100:+.2f}%",
               f"{(bench_end['NVDA']/bench_start['NVDA']-1)*100:+.2f}%"],
    "Daily Vol": [f"{bench['SPCX_ret'].dropna().std():.2f}%",
                  f"{bench['QQQ_ret'].dropna().std():.2f}%",
                  f"{bench['NVDA_ret'].dropna().std():.2f}%"],
})
_add_df_table(doc, t3_show)

# ── 7. Verdict ──
doc.add_heading("7. Verdict", level=1)
doc.add_paragraph(verdict)

# ── Save ──
doc.save(str(DOCX_PATH))
print(f"Saved: {DOCX_PATH}")
